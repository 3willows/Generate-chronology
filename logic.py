import re
import csv
import pandas as pd
import os
import pypandoc

from docx import Document
from flask import Flask

from pathlib import Path 
from operator import itemgetter

from utils import parse_date

app = Flask(__name__)

def everything_function(f: Path) -> None:
    
    # Define the input and output paths
    steps = [
            (convert_to_markdown, [f, Path('input.md')]),
            (clean_markdown_document, [Path('input.md'), Path('cleaned.md')]),
            (extract_numbered_items_to_csv, [Path('cleaned.md'), Path('all_dates.csv')]),
            (extract_dates, [Path('all_dates.csv'), Path('dates_extracted.csv')]),
            (create_word_document_from_csv, [Path('dates_extracted.csv'), Path('draft-chronology.docx')])
        ]
    intermediate_files = []

    try:
        for step, args in steps:
            step(*args)

    finally:
           for file in [f, 'input.md', 'cleaned.md', 'all_dates.csv', 'dates_extracted.csv']:
            if os.path.exists(file):
                os.remove(file)
                print(f"Removed {file}")
 
def convert_to_markdown(input_file: Path, output_file: Path) -> Path:
    pypandoc.convert_file(input_file, 'md', outputfile=str(output_file))
    return output_file

def clean_markdown_document(input_file: Path, output_file: Path) -> None:
    with input_file.open('r', encoding='utf-8') as file:
        content = file.read()

    first_occurrence_index = content.find("1.")
    first_occurrence_of_footnote = content.find("[^1]:")  

    if first_occurrence_index != -1:
        if first_occurrence_of_footnote != -1 and first_occurrence_of_footnote > first_occurrence_index:
            cleaned_content = content[:first_occurrence_of_footnote]
        else:
            cleaned_content = content[first_occurrence_index:]
    else:
        cleaned_content = content

    with output_file.open('w', encoding='utf-8') as file:
        file.write(cleaned_content)

    print(f"Markdown document cleaned and saved as {output_file}")

def extract_numbered_items_to_csv(input_file: Path, output_file: Path) -> None:
    with input_file.open('r', encoding='utf-8') as file:
        content = file.read()

    pattern = r'(\d+)\.\s*(.*?)\n(?=\d+\.\s|$)'
    matches = re.findall(pattern, content, re.DOTALL)

    with output_file.open('w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Number', 'Text'])

        for number, text in matches:
            writer.writerow([number, text.strip()])

    print(f"Data has been extracted and saved to {output_file}")

def extract_dates(input_file: Path, output_file: Path) -> None:
    extracted_data = []
    date_patterns = [
        r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b',
        r'\b(\d{1,2} [A-Za-z]{3,9} \d{2,4})\b',
        r'\b([A-Za-z]{3,9} \d{1,2}, \d{2,4})\b',
        r'\b(\d{4}-\d{2}-\d{2})\b'
    ]

    with input_file.open('r', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        
        for row in reader:
            paragraph_number = row['Number']
            text = row['Text']
            
            for pattern in date_patterns:
                matches = re.findall(pattern, text)
                for match in matches:
                    date = parse_date(match)
                    if date:
                        extracted_data.append({'Date': date, 'Text': text, 'Paragraph Number': paragraph_number})

    extracted_data.sort(key=itemgetter('Date'))

    with output_file.open('w', newline='', encoding='utf-8') as csvfile:
        fieldnames = ['Date', 'Text', 'Paragraph Number']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        
        writer.writeheader()
        
        for entry in extracted_data:
            writer.writerow({'Date': entry['Date'].strftime("%Y-%m-%d"), 
                             'Text': entry['Text'], 
                             'Paragraph Number': entry['Paragraph Number']})

    print(f"Date extraction completed and saved to {output_file}")

def create_word_document_from_csv(input_file: Path, output_file: Path) -> None:
    data = pd.read_csv(input_file)

    doc = Document()
    doc.add_heading('Draft Chronology', level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Text'
    hdr_cells[2].text = 'Paragraph Number'

    for index, row in data.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Date'])
        row_cells[1].text = str(row['Text'])
        row_cells[2].text = str(row['Paragraph Number'])

    doc.save(output_file)

    print(f"Word document '{output_file}' created successfully.")

@app.route('/')   
def main():
    everything_function(Path('your_input_file_here')) 

if __name__ == '__main__':   
    app.run(debug=True)